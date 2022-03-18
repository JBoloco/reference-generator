import streamlit as st
import docx

st.set_page_config(layout="centered", page_icon="📚", page_title="Reference Generator")
st.title("📚 Reference Generator")

st.write(
    "This app will format a reference for a book, article or website"
)

with st.expander(" ℹ️ - About this app", expanded=False):

    st.write(
        """
- This *Reference Generator* app aims to make formatting citations for a bibliography easier following a Harvard style referencing.
- It currently has no option to copy to clipboard but has an integrated option to output the reference to a new word document.

        """

    )

col1, col2 = st.columns([2, 1.5])

col2.write("  ")
ref = col1.radio(
    "Choose what you would like to reference",
    ["Book", "Article", "Website"],
    help="The 3 options will give 3 different formats",
    index=0,
)

def create_book():

    with st.form("my_form"):

        submit = st.form_submit_button("Submit")
        error = IndexError("Please enter the author's names")
        
        doc = col1.checkbox(
            "Tick if you would like to print the reference to a word document",
            help="The reference will be pasted into a new word document",
        )
        
        authors = col1.radio(
            "Select the number of authors",
            ["1", "2", "More than 2 authors"],
            help="If there are more than 2 authors then the reference will automatically add 'et al.'",
            index=0,
            disabled=False,

        )
        
        if authors == "More than 2 authors":
            forename = col1.text_input("Author's first name")
            surname = col1.text_input("Author's surname")
            year = col2.text_input("Year of publishing")
            book = col2.text_input("Title of the book")
            edition = col2.text_input("Edition of the book (e.g 2nd)")
            publication = col2.text_input("Place of publication")
            publisher = col2.text_input("Publisher")

            st.subheader("Your Reference:")

            if doc:
                doc_name = col2.text_input("What would you like to name your document? (Do not include the '.docx' file extension)")
                doc = docx.Document()
                doc.add_paragraph(f'\n{surname.title()}, {forename[0].title()}. et al. ({year}). "{book.title()}." {edition} Edition. {publication.title()}: {publisher}.')
                doc.save(f"{doc_name}.docx")

            if submit:
                try:
                    st.markdown(f'{surname.title()}, {forename[0].title()}. et al. ({year}). *{book.title()}*. {edition} Edition. {publication.title()}: {publisher}.')
                except IndexError:
                    st.exception(error)

        elif authors == "2":
            forename = col1.text_input("Author's first name")
            surname = col1.text_input("Author's surname")
            forename2 = col1.text_input("Second author's first name")
            surname2 = col2.text_input("Second author's surname")
            year = col2.text_input("Year of publishing")
            book = col2.text_input("Title of the book")
            edition = col2.text_input("Edition of the book (e.g 2nd)")
            publication = col2.text_input("Place of publication")
            publisher = col2.text_input("Publisher")

            st.subheader("Your Reference:")

            if doc:
                doc_name = col2.text_input("What would you like to name your document? (Do not include the '.docx' file extension)")
                doc = docx.Document()
                doc.add_paragraph(f'\n{surname.title()}, {forename[0].title()}. and {surname2.title()}, {forename2[0].title()}. ({year}). "{book.title()}" {edition} Edition. {publication.title()}: {publisher}.')
                doc.save(f"{doc_name}.docx")

            if submit:
                try:
                    st.markdown(f'{surname.title()}, {forename[0].title()}. and {surname2.title()}, {forename2[0].title()}. ({year}). *{book.title()}* {edition} Edition. {publication.title()}: {publisher}.')
                except IndexError:
                    st.exception(error)
            
        elif authors == "1":
            forename = col1.text_input("Author's first name")
            surname = col1.text_input("Author's surname")
            year = col2.text_input("Year of publishing")
            book = col2.text_input("Title of the book")
            edition = col2.text_input("Edition of the book (e.g 2nd)")
            publication = col2.text_input("Place of publication")
            publisher = col2.text_input("Publisher")

            st.subheader("Your Reference:")
            
            if doc:
                doc_name = col2.text_input("What would you like to name your document? (Do not include the '.docx' file extension)")
                doc = docx.Document()            
                doc.add_paragraph(f'\n{surname.title()}, {forename[0].title()}. ({year}). "{book.title()}." {edition} Edition. {publication.title()}: {publisher}.')
                doc.save(f"{doc_name}.docx")

            if submit:
                try:
                    st.markdown(f'{surname.title()}, {forename[0].title()}. ({year}). *{book.title()}*. {edition} Edition. {publication.title()}: {publisher}.')
                except IndexError:
                    st.exception(error)

def create_article():
    """Formatting conditions for the number of authors in the reference"""

    with st.form("my_form"):

        submit = st.form_submit_button("Submit")
        error = IndexError("Please enter the author's names")
        
        doc = col1.checkbox(
            "Tick if you would like to print the reference to a word document",
            help="The reference will be pasted into a new word document",
        )

        authors = col1.radio(
            "Select the number of authors",
            ["1", "2", "More than 2 authors"],
            help="If there are more than 2 authors then the reference will automatically add 'et al.'",
            index=0,
            disabled=False,

        )
        
        if authors == "More than 2 authors":
            forename = col1.text_input("Author's first name")
            surname = col1.text_input("Author's surname")
            year = col1.text_input("Year of publication")
            article = col2.text_input("Title of the article")
            journal = col2.text_input("Journal where the article is found")
            volume = col2.text_input("Journal volume")
            issue = col2.text_input("Journal issue")
            pages = col2.text_input("Range of pages cited (e.g. 1-5)")

            st.subheader("Your Reference:")

            if doc:
                doc_name = col2.text_input("What would you like to name your document? (Do not include the '.docx' file extension)")
                doc = docx.Document()
                doc.add_paragraph(f'\n{surname.title()}, {forename[0].title()}. et al. ({year}) "{article.title()}" {journal.title()}, {volume}({issue}): pp.{pages}.')
                doc.save(f"{doc_name}.docx")

            if submit:
                try:
                    st.markdown(f"{surname.title()}, {forename[0].title()}. et al. ({year}) '{article.title()}' *{journal.title()}*, {volume}({issue}): pp.{pages}.")
                except IndexError:
                    st.exception(error)
            
        elif authors == "2":
            forename = col1.text_input("Author's first name")
            surname = col1.text_input("Author's surname")
            forename2 = col1.text_input("Second author's first name")
            surname2 = col1.text_input("Second author's surname")
            year = col2.text_input("Year of publication")
            article = col2.text_input("Title of the article")
            journal = col2.text_input("Journal where the article is found")
            volume = col2.text_input("Journal volume")
            issue = col2.text_input("Journal issue")
            pages = col2.text_input("Range of pages cited (e.g. 1-5)")

            st.subheader("Your Reference:")

            if doc:
                doc_name = col2.text_input("What would you like to name your document? (Do not include the '.docx' file extension)")
                doc = docx.Document()
                doc.add_paragraph(f"{surname.title()}, {forename[0].title()}. and {surname2.title()}, {forename2[0].title()}. ({year}) '{article.title()}' {journal.title()}, {volume}({issue}): pp.{pages}.")
                doc.save(f"{doc_name}.docx")

            if submit:
                try:
                    st.markdown(f"{surname.title()}, {forename[0].title()}. and {surname2.title()}, {forename2[0].title()}. ({year}) '{article.title()}' *{journal.title()}*, {volume}({issue}): pp.{pages}.")
                except IndexError:
                    st.exception(error)
      
        elif authors == "1":
            forename = col1.text_input("Author's first name")
            surname = col1.text_input("Author's surname")
            year = col1.text_input("Year of publication")
            article = col2.text_input("Title of the article")
            journal = col2.text_input("Journal where the article is found")
            volume = col2.text_input("Journal volume")
            issue = col2.text_input("Journal issue")
            pages = col2.text_input("Range of pages cited (e.g. 1-5)")

            st.subheader("Your Reference:")
            
            if submit:
                try:
                    st.markdown(f"{surname.title()}, {forename[0].title()}. ({year}) '{article.title()}' *{journal.title()}*, {volume}({issue}): pp.{pages}.")
                except IndexError:
                    st.exception(error)
            if doc:
                doc_name = col2.text_input("What would you like to name your document? (Do not include the '.docx' file extension)")
                doc = docx.Document()
                doc.add_paragraph(f"\n{surname.title()}, {forename[0].title()}. ({year}) '{article.title()}' {journal.title()}, {volume}({issue}): {pages}.")
                doc.save(f"{doc_name}.docx")

def create_web():
    """Format a reference from a website"""
        
    with st.form("my_form"):

        submit = st.form_submit_button("Submit")
        error = IndexError("Please enter the author's names")
        org_error = ValueError("Please enter the name of the organisation")
        
        doc = col1.checkbox(
            "Tick if you would like to print the reference to a new word document",
            help="The reference will be pasted into a new word document each time",
        )

        authors = col1.radio(
        "Select whether author or organisation",
        ["Author", "Group/Organisation"],
        help="If the web page has no particular author simply enter the name of the group/organisation",
        index=0,
        disabled=False,

        )
        if authors == "Author":
            forename = col1.text_input("Author/Editor's first name")
            surname = col1.text_input("Author/Editor's surname")
            year = col2.text_input("Year of publication")
            name = col2.text_input("Title and subtitle of the online article")
            link = col2.text_input("Enter the URL")
            date = col2.date_input("Date of access")
            
            if doc:
                doc_name = col2.text_input("What would you like to name your new document? (Do not include the '.docx' file extension)")
                doc = docx.Document()
                doc.add_paragraph(f'\n{surname.title()}, {forename[0].title()}. ({year}) "{name.title()}". Available at: {link} [Accessed on: {date}]')
                doc.save(f"{doc_name}.docx")

            st.subheader("Your Reference:")
        
            if submit:
                try:
                    st.markdown(f'{surname.title()}, {forename[0].title()}. ({year}) "*{name.title()}*". Available at: {link} [Accessed on: {date}]')
                except IndexError:
                    st.exception(error)

        elif authors == "Group/Organisation":
            org = col1.text_input("Name of the organisation")
            year = col2.text_input("Year of publication")
            name = col2.text_input("Title and subtitle of the online article")
            link = col2.text_input("Enter the URL")
            date = col2.date_input("Date of access")
            
            if doc:
                doc_name = col2.text_input("What would you like to name your new document? (Do not include the '.docx' file extension)")
                doc = docx.Document()
                doc.add_paragraph(f'\n{org}. ({year}) "{name.title()}". Available at: {link} [Accessed on: {date}]')
                doc.save(f"{doc_name}.docx")

            st.subheader("Your Reference:")
        
            if submit:
                if len(org)==0:
                    st.exception(org_error)
                else:
                    st.markdown(f'{org}. ({year}) "*{name.title()}*". Available at: {link} [Accessed on: {date}]')


def choice():
    """Select which reference"""
    
    if ref == "Book":
        create_book()
        
    elif ref == "Article":
        create_article()
        
    elif ref == "Website":
        create_web()

choice()
